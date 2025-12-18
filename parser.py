import os
from pathlib import Path
from docling.document_converter import DocumentConverter, PdfFormatOption
from docling.datamodel.pipeline_options import PdfPipelineOptions
from docling.datamodel.base_models import InputFormat


def extract_docx_metadata_docx2python(file_path: str) -> str:
   
    try:
        from docx2python import docx2python
        
        # Extract all content including headers, footers, and text boxes
        with docx2python(file_path) as doc:
            extracted_text = []
            
            # Extract header content (nested list structure)
            if doc.header:
                header_text = _flatten_docx2python_content(doc.header)
                for text in header_text:
                    if text and text not in extracted_text:
                        extracted_text.append(text)
            
            # Extract footer content (this is where CRICOS codes often hide)
            if doc.footer:
                footer_text = _flatten_docx2python_content(doc.footer)
                for text in footer_text:
                    if text and text not in extracted_text:
                        extracted_text.append(text)
            
            # Extract text from footnotes (sometimes metadata is here)
            if doc.footnotes:
                footnote_text = _flatten_docx2python_content(doc.footnotes)
                for text in footnote_text:
                    if text and text not in extracted_text:
                        extracted_text.append(text)
        
        if extracted_text:
            return "--- DOCUMENT HEADER/FOOTER METADATA ---\n" + "\n".join(extracted_text) + "\n--- END METADATA ---\n\n"
        return ""
        
    except Exception as e:
        print(f"Warning: docx2python extraction failed: {e}")
        # Fallback to python-docx method
        return _extract_docx_headers_footers_fallback(file_path)


def _flatten_docx2python_content(content) -> list:

    result = []
    
    def recurse(item):
        if isinstance(item, str):
            text = item.strip()
            if text:
                result.append(text)
        elif isinstance(item, list):
            for sub_item in item:
                recurse(sub_item)
    
    recurse(content)
    return result


def _extract_docx_headers_footers_fallback(file_path: str) -> str:
    """
    Fallback method using python-docx if docx2python fails.
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        extracted_text = []
        
        for section in doc.sections:
            # Extract header text
            header = section.header
            if header:
                for paragraph in header.paragraphs:
                    text = paragraph.text.strip()
                    if text and text not in extracted_text:
                        extracted_text.append(text)
            
            # Extract footer text
            footer = section.footer
            if footer:
                for paragraph in footer.paragraphs:
                    text = paragraph.text.strip()
                    if text and text not in extracted_text:
                        extracted_text.append(text)
        
        if extracted_text:
            return "--- DOCUMENT HEADER/FOOTER METADATA ---\n" + "\n".join(extracted_text) + "\n--- END METADATA ---\n\n"
        return ""
        
    except Exception as e:
        print(f"Warning: Fallback extraction also failed: {e}")
        return ""


def convert_to_markdown(file_path: str) -> str:

    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")
    
    suffix = path.suffix.lower()
    if suffix not in [".pdf", ".docx"]:
        raise ValueError(f"Unsupported file format: {suffix}. Use PDF or DOCX.")
    
    # Extract header/footer metadata for DOCX files using docx2python
    header_footer_text = ""
    if suffix == ".docx":
        header_footer_text = extract_docx_metadata_docx2python(file_path)
    
    # Configure pipeline for fast programmatic parsing (no OCR)
    pipeline_options = PdfPipelineOptions()
    pipeline_options.do_ocr = False  # Disable OCR for digital documents
    pipeline_options.do_table_structure = True  # Preserve table structure
    
    converter = DocumentConverter(
        format_options={
            InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
        }
    )
    
    # Convert document
    result = converter.convert(file_path)
    
    # Export to Markdown
    markdown_content = result.document.export_to_markdown()
    
    # Prepend header/footer metadata so LLM can extract it
    return header_footer_text + markdown_content


if __name__ == "__main__":
    # Quick test
    import sys
    if len(sys.argv) > 1:
        md = convert_to_markdown(sys.argv[1])
        print(md)

